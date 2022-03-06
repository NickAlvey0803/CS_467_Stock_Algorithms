# Initial Network Testing
# Written by Nick Alvey

#####################################################################################################################
# Import Necessary Components
#####################################################################################################################

# For Data Parsing
import docx
from docx2pdf import convert
import pandas as pd
import glob
import os
import numpy as np
from scipy.stats import linregress
import datetime
import matplotlib.pyplot as plt
import sys
# import docx
# from docx2pdf import convert
from keras.callbacks import Callback
# For Neural Networks

import tensorflow as tf
from tensorflow import keras
from tensorflow.keras import layers
from keras.models import load_model

# Add access to the data folder
sys.path.insert(1, '../Data')

#####################################################################################################################

# Begin Initialization of Algorithm

#####################################################################################################################

# List CPU's Available
print("------------------------------------------------------------")
print("Num CPUs Available: ", len(tf.config.list_physical_devices('CPU')))
print("------------------------------------------------------------")
print("Run with Tensorflow Version: " + tf.__version__)
print("------------------------------------------------------------")


#####################################################################################################################

# Get CSV for Volatility and Momentum Calculations
# This can be over any number of ETF's, or just a single one

#####################################################################################################################

def CreateDataset(name_to_return, lead_up_days=2, consideration_days=90, putcallflag=0,
                  junkbondflag=0, mcclellanflag=0, future_num_days=30):
    # Number of days in sample to consider leading up to current day, lowest it should be is 2
    ETFfile = '../Data/3x-ETF/' + name_to_return + '.csv'
    N = lead_up_days
    """

    :param ETFpath:
    :param N:
    :param consideration_days:
    :param putcallflag:
    :param junkbondflag:
    :param mcclellanflag:
    :param future_num_days:
    :return: The compiled ETF dataset


    """
    Whole_ETF_3x = None

    # ETFpath = '../Data/3x-ETF/'  # Reference '3x-ETF-All-Files/' to run over all 3x ETF's

    # Flag to initialize Neural Network input
    first_run_flag = 0
    name = str(name_to_return) + ".csv"
    print_and_write_status("\n\nBuilding Dataset for: " + str(name_to_return))
    ETF_3x = pd.read_csv(ETFfile)

    ############################################################################################################

    # Price - Derived Calculations

    ############################################################################################################

    # Calculating Historic Volatility from Raw 3x ETF Data

    # Theory Yang-Zhang (OHLC)
    # Sum of overnight volatility and weighted average of Rogers-Satchell volatility
    # NOTE: In the Paper: MEASURING HISTORICAL VOLATILITY, temp_overnight_vol and temp_openclose_vol if
    # implemented exactly would be 0 every time - it subtracts one value from its equal. Implementations of this
    # algorithm do not do that, so I did not either.

    # Number of days in sample to consider leading up to current day, lowest it should be is 2
    # N = 2

    # Scaling factor
    F = 1

    k = 0.34 / (1.34 + (N + 1) / (N - 1))

    yang_zhang_vol_list = [np.nan] * N

    for day in range(N, len(ETF_3x)):
        temp_rog_satch_vol = 0  # initialize to zero for every consideration
        temp_overnight_vol = 0  # initialize to zero for every consideration
        temp_openclose_vol = 0  # initialize to zero for every consideration
        for i in range(N):
            temp_rog_satch_vol += np.log(ETF_3x['High'][day - i] / ETF_3x['Close'][day - i]) * np.log(
                ETF_3x['High'][day - i] / ETF_3x['Open'][day - i]) + np.log(
                ETF_3x['Low'][day - i] / ETF_3x['Close'][day - i]) * np.log(
                ETF_3x['Low'][day - i] / ETF_3x['Open'][day - i])
            temp_overnight_vol += (np.log(
                ETF_3x['Open'][day - i] / ETF_3x['Close'][day - np.absolute((i - 1))])) ** 2
            temp_openclose_vol += (np.log(ETF_3x['Close'][day - i] / ETF_3x['Open'][day - i])) ** 2
        rog_satch_vol = temp_rog_satch_vol * F / N
        overnight_vol = temp_overnight_vol * F / (N - 1)
        openclose_vol = temp_openclose_vol * F / (N - 1)

        yang_zhang_vol = np.sqrt(F) * np.sqrt(overnight_vol + k * openclose_vol + (1 - k) * rog_satch_vol)

        yang_zhang_vol_list.append(yang_zhang_vol)

    ETF_3x['Volatility'] = yang_zhang_vol_list

    # print_and_write_status(name + " Volatility Calculated and Appended to Dataframe")
    # print_and_write_status("------------------------------------------------------------")

    ############################################################################################################

    # Calculating Momentum from Raw 3x ETF Data

    # Used from Tutorial here: https://teddykoker.com/2019/05/momentum-strategy-from-stocks-on-the-move-in-python/
    # Theory from Andreas Clenow

    # Number of days to calculate momentum over - it is not recommended to go below 30, and
    # 90 is recommended by link. This number changes the output drastically.

    if len(ETF_3x) < consideration_days:
        consideration_days = len(ETF_3x)
    else:
        # consideration_days = 90
        pass

    momentum_list = [np.nan] * consideration_days

    for days in range(consideration_days, len(ETF_3x)):
        consideration_days_list = []
        for datapoint in range(consideration_days):
            consideration_days_list.append(ETF_3x['Close'][days - datapoint])
        returns = np.log(consideration_days_list)
        x = np.arange(len(consideration_days_list))
        slope, _, rvalue, _, _ = linregress(x, returns)

        momentum = (1 + slope) ** 252 * (rvalue ** 2)

        momentum_list.append(momentum)

    ETF_3x['Momentum'] = momentum_list

    # print_and_write_status(name + " " + str(consideration_days) + "-day Momentum Calculated and Appended to Dataframe")
    # print_and_write_status("------------------------------------------------------------")

    ############################################################################################################

    # Sentimental Factors

    ############################################################################################################

    # Compiling Put/Call Ratios
    if putcallflag == 0:
        PutCallpath = '../Data/Put-Call-Ratio/'
        PutCallfile = 'totalpc.csv'

        PutCall_rawdata = pd.read_csv(PutCallpath + PutCallfile)

        ETF_3x['Put/Call Ratio'] = np.nan

        for days in range(len(ETF_3x)):
            etf_datetime = datetime.datetime.strptime(ETF_3x['Date'][days], "%Y-%m-%d")
            putcall_index = PutCall_rawdata[
                PutCall_rawdata['date'] == etf_datetime.strftime('%Y-%m-%d')].index.values
            if len(PutCall_rawdata.iloc[putcall_index, PutCall_rawdata.columns.get_loc('p_c_ratio')]) == 0:
                ETF_3x.iloc[days, ETF_3x.columns.get_loc('Put/Call Ratio')] = np.nan
            else:
                ETF_3x.iloc[days, ETF_3x.columns.get_loc('Put/Call Ratio')] = \
                    PutCall_rawdata.iloc[putcall_index, PutCall_rawdata.columns.get_loc('p_c_ratio')]

        # print_and_write_status(name + " Put/Call Ratio Compiled and Appended to Dataframe")
        # print_and_write_status("------------------------------------------------------------")
    else:
        pass

    ############################################################################################################

    # Calculating Junk Bond Demand (Volume)

    # Indicator for down days or up days not implemented
    if junkbondflag == 0:
        JunkBondpath = '../Data/Junk-Bond-ETF/'  # Should be updated to eventual folder name
        JunkBondfiles = glob.glob(os.path.join(JunkBondpath, "*.csv"))

        frames = []
        for junk_bonds_csvs in JunkBondfiles:
            temp_JunkBond = pd.read_csv(junk_bonds_csvs)
            frames.append(temp_JunkBond["Date"])
            frames.append(temp_JunkBond["Volume"])

        headers = ["Date1", "Volume1", "Date2", "Volume2", "Date3", "Volume3", "Date4", "Volume4", "Date5",
                   "Volume5"]  # hardcoded since requirement is "Top 5 ETF's to be Considered"
        JunkBond_rawdata = pd.concat(frames, axis=1, keys=headers)
        JunkBond_rawdata['Total Volume'] = np.nan

        junkbond_datetime_start = datetime.date(2007, 1, 1)
        junkbond_datetime_end = datetime.date.today()

        days = junkbond_datetime_start
        listpair = []

        while days < junkbond_datetime_end:
            daily_volume = 0

            if len(JunkBond_rawdata[JunkBond_rawdata['Date1'] == days.strftime('%Y-%m-%d')].index.tolist()) > 0:
                index = JunkBond_rawdata[JunkBond_rawdata['Date1'] == days.strftime('%Y-%m-%d')].index.tolist()
                daily_volume += JunkBond_rawdata['Volume1'][index[0]]
            if len(JunkBond_rawdata[JunkBond_rawdata['Date2'] == days.strftime('%Y-%m-%d')].index.tolist()) > 0:
                index = JunkBond_rawdata[JunkBond_rawdata['Date2'] == days.strftime('%Y-%m-%d')].index.tolist()
                daily_volume += JunkBond_rawdata['Volume2'][index[0]]
            if len(JunkBond_rawdata[JunkBond_rawdata['Date3'] == days.strftime('%Y-%m-%d')].index.tolist()) > 0:
                index = JunkBond_rawdata[JunkBond_rawdata['Date3'] == days.strftime('%Y-%m-%d')].index.tolist()
                daily_volume += JunkBond_rawdata['Volume3'][index[0]]
            if len(JunkBond_rawdata[JunkBond_rawdata['Date4'] == days.strftime('%Y-%m-%d')].index.tolist()) > 0:
                index = JunkBond_rawdata[JunkBond_rawdata['Date4'] == days.strftime('%Y-%m-%d')].index.tolist()
                daily_volume += JunkBond_rawdata['Volume4'][index[0]]
            if len(JunkBond_rawdata[JunkBond_rawdata['Date5'] == days.strftime('%Y-%m-%d')].index.tolist()) > 0:
                index = JunkBond_rawdata[JunkBond_rawdata['Date5'] == days.strftime('%Y-%m-%d')].index.tolist()
                daily_volume += JunkBond_rawdata['Volume5'][index[0]]

            listpair.append([days.strftime('%Y-%m-%d'), daily_volume])

            days = days + datetime.timedelta(days=1)

        JunkBond_interpreted_data = pd.DataFrame(listpair, columns=['Date', 'Volume'])

        ETF_3x['Junk Bond Demand'] = np.nan

        for days in range(len(ETF_3x)):
            etf_datetime = datetime.datetime.strptime(ETF_3x['Date'][days], "%Y-%m-%d")

            junkbond_index = JunkBond_interpreted_data[
                JunkBond_interpreted_data['Date'] == etf_datetime.strftime('%Y-%m-%d')].index.values

            ETF_3x.iloc[days, ETF_3x.columns.get_loc('Junk Bond Demand')] = \
                JunkBond_interpreted_data.iloc[
                    junkbond_index, JunkBond_interpreted_data.columns.get_loc('Volume')]

        # print_and_write_status(name + " Junk Bond Demand Calculated and Appended to Dataframe")
        # print_and_write_status("------------------------------------------------------------")
    else:
        pass

    ############################################################################################################

    # Calculating McClellan Summation Index on stock in question

    # Sources:
    # https://www.investopedia.com/terms/m/mcclellansummation.asp
    # https://www.investopedia.com/terms/m/mcclellanoscillator.asp
    # https://www.investopedia.com/terms/e/ema.asp

    # Initialize temporary variables
    if mcclellanflag == 0:
        advance = []
        EMA19 = []
        EMA39 = []
        SMA19 = 0
        SMA39 = 0

        MCsummation_index = 0
        ETF_3x['McClellan Summation Index'] = np.nan

        # Requires Date to be structured Old --> New
        # Adjusted for Size (normalized per Adjusted Oscillator formula)

        for index in range(len(ETF_3x)):

            # Calculate Advances and Declines for each day
            advance.append((ETF_3x['Close'][index] - ETF_3x['Open'][index]) /
                           (ETF_3x['Close'][index] + ETF_3x['Open'][index]))

            if index == 17:
                SMA19 = np.average(advance[0:index])

            elif index == 18:
                EMA19.append((advance[index]) - SMA19 * 0.1 + SMA19)

            elif index > 18:
                EMA19.append((advance[index]) - EMA19[index - 19] * 0.1 + EMA19[index - 19])

                if index == 37:
                    SMA39 = np.average(advance[0:index])

                elif index == 38:
                    EMA39.append((advance[index]) - SMA39 * 0.05 + SMA19)

                elif index > 38:
                    EMA39.append((advance[index]) - EMA39[index - 39] * 0.05 + EMA39[index - 39])

                    # Convert to McClellan Oscillator

                    adjusted_MCOscillator = EMA19[index - 19] - EMA39[index - 39]

                    # Convert to McClellan Summation Index

                    MCsummation_index = MCsummation_index + adjusted_MCOscillator

                    ETF_3x.iloc[index, ETF_3x.columns.get_loc('McClellan Summation Index')] = MCsummation_index

        # print_and_write_status(name + " McClellan Summation Index Calculated and Appended to Dataframe")
        # print_and_write_status("------------------------------------------------------------")
    else:
        pass

    ############################################################################################################

    # Add in Profit for Traded Data via QuantConnect Backtesting of Provided Algorithm

    ############################################################################################################

    Quantpath = '../Data/3X-ETF-Backtests-CSharp/'

    Quantfile = name
    # print_and_write_status("Name is: ",name)
    Quant_rawdata = None

    # Note, for now, columns are off by one, so "Date" here equates to the column of "Value"

    skip_flag = 0
    try:
        Quant_rawdata = pd.read_csv(Quantpath + name)
    except:
        print_and_write_status(
            "ALERT: Profit Data not Formatted in the Same Way as Rest of Profit Data, Please Check")
        print_and_write_status("This ETF will be Ignored and Removed from Dataset")
        print_and_write_status("------------------------------------------------------------")
        skip_flag = 1

    if skip_flag == 0:
        Quant_dates = Quant_rawdata['date'].tolist()
        ETF_3x['Profit Percentage'] = np.nan

        start = 0  # Assuming starting money amount is $100,000, should be entered by user
        temp_percent = 0

        for days in range(len(ETF_3x)):
            flag = 0
            temp = 0
            for dates in Quant_dates:
                if ETF_3x['Date'][days] == dates:
                    fill_loc = Quant_rawdata.index[Quant_rawdata['date'] == dates].tolist()
                    if len(fill_loc) > 1:
                        fill_loc = fill_loc[1]
                    if float(Quant_rawdata['fill-quantity'][fill_loc]) != 0:
                        temp += float(Quant_rawdata['fill-quantity'][fill_loc])

            if start == 0:
                temp_percent = 1

            if temp == 0:
                pass

            else:
                if start == 0:
                    temp_percent = 1
                else:
                    temp_percent = temp / start * 100

            start += temp

            # do temp_percent for percentage change, start for total amount in inventory
            ETF_3x.iloc[days, ETF_3x.columns.get_loc('Profit Percentage')] = temp_percent

        # print_and_write_status(name + " Profit Percentage by Trading Algorithm Compiled and Appended to Dataframe")
        # print_and_write_status("------------------------------------------------------------")
    else:
        ETF_3x['Profit Percentage'] = np.nan
    ############################################################################################################

    # Add in Volatility Time Lag Variables

    ############################################################################################################

    num_days = 10  # starting value, can be changed to any number

    delta_between = 1  # starting value, can be changed to any number

    for n in range(num_days // delta_between):  # must be whole number
        ETF_3x['Volatility Time Lag ' + str(n)] = np.nan

        for days in range(len(ETF_3x)):
            if days >= n * delta_between:
                ETF_3x.iloc[days, ETF_3x.columns.get_loc('Volatility Time Lag ' + str(n))] = \
                    ETF_3x.iloc[days - n * delta_between, ETF_3x.columns.get_loc('Volatility')]

    # print_and_write_status(name + " " + str(num_days) + "-day Volatility Lag with " + str(
    # delta_between) + "-day Spacing Compiled and "
    # "Appended to Dataframe")
    # print_and_write_status("------------------------------------------------------------")

    ############################################################################################################

    # Copy ETF for last 30 days for prediction purposes

    ############################################################################################################

    prediction_ETF = ETF_3x.dropna()
    prediction_ETF = prediction_ETF.tail(30)
    prediction_ETF = prediction_ETF.drop(columns=['Date'])
    # prediction_ETF.to_csv('../Prediction/' + 'prediction_' + name_to_return)

    ############################################################################################################

    # Calculate Future Profit Prediction

    ############################################################################################################

    # future_num_days = 30  # starting value, can be changed to any number

    ETF_3x['Profit Prediction' + str(future_num_days) + ' days from now'] = np.nan

    for days in range(len(ETF_3x) - (future_num_days + 1)):
        ETF_3x.iloc[
            days, ETF_3x.columns.get_loc('Profit Prediction' + str(future_num_days) + ' days from now')] = 0
        for n in range(future_num_days):
            ETF_3x.iloc[days, ETF_3x.columns.get_loc(
                'Profit Prediction' + str(future_num_days) + ' days from now')] += \
                ETF_3x.iloc[days + n + 1, ETF_3x.columns.get_loc('Profit Percentage')]

        ETF_3x.iloc[
            days, ETF_3x.columns.get_loc('Profit Prediction' + str(future_num_days) + ' days from now')] \
            = (ETF_3x.iloc[days, ETF_3x.columns.get_loc('Profit Prediction' + str(future_num_days)
                                                        + ' days from now')] / future_num_days)

    # print_and_write_status(name + " " + str(
    # future_num_days) + "-day Delta in Profit Made from the Trading Algorithm Compiled and Appended to Dataframe")
    # print_and_write_status("------------------------------------------------------------")

    ############################################################################################################

    ############################################################################################################

    # Some Last Cleanup of the Dataframe

    ############################################################################################################

    # If desired, Turn Date into Separate Column for Year, Month, and Date

    # ETF_3x['Year'] = np.nan
    # ETF_3x['Month'] = np.nan
    # ETF_3x['Day'] = np.nan

    # for days in range(len(ETF_3x)):
    #     etf_datetime = datetime.datetime.strptime(ETF_3x['Date'][days], "%Y-%m-%d")
    #     ETF_3x.iloc[days, ETF_3x.columns.get_loc('Year')] = etf_datetime.strftime('%Y')
    #     ETF_3x.iloc[days, ETF_3x.columns.get_loc('Month')] = etf_datetime.strftime('%m')
    #     ETF_3x.iloc[days, ETF_3x.columns.get_loc('Day')] = etf_datetime.strftime('%d')

    ############################################################################################################

    # Drop Data that isn't valuable, for now that's just the "Data" Column because it's not convertible to a float
    # If it is desired, uncomment the section above

    ETF_3x = ETF_3x.drop(columns=['Date'])

    ############################################################################################################

    # Removes all Nan values
    ETF_3x = ETF_3x.dropna()
    ETF_3x = ETF_3x.reset_index(drop=True)

    # if skip_flag == 1:
    # print_and_write_status("This ETF is being removed from the Dataset")
    # print_and_write_status("------------------------------------------------------------")
    # else:
    # print_and_write_status("The ETF Dataset for " + name + " has been Finalized and is Being Added to the Overall Data Set")
    # print_and_write_status("------------------------------------------------------------")

    if first_run_flag == 0:
        Whole_ETF_3x = ETF_3x
        first_run_flag = 1
    else:
        Whole_ETF_3x = pd.concat([Whole_ETF_3x, ETF_3x], ignore_index=True)

    ################################################################################################################

    # Some Troubleshooting Prints - Printing out the shape of the Dataframe, as well as outputting it to a CSV
    # To view potential errors
    #
    # print_and_write_status("------------------------------------------------------------")
    # print_and_write_status("SHAPE IS: ")
    # print_and_write_status(Whole_ETF_3x.shape)
    # print_and_write_status("------------------------------------------------------------")
    print_and_write_status("Saving dataset...")
    Whole_ETF_3x.to_csv('../Data/Built-Datasets/' + str(tfd['Model_Name']))

    return Whole_ETF_3x, future_num_days, name_to_return, prediction_ETF


# future_num_days = 15  # starting value, can be changed to any number

#####################################################################################################################

# Preparing to Run Neural Network

#####################################################################################################################

def CreateNeuralNetwork(Whole_ETF_3x, prediction_ETF, name_to_return, future_num_days=30, models_to_test=30, lim1=10,
                        lim2=18, lim3=24, base_epochs=500, base_learning_rate=0.001):
    """

    :param Whole_ETF_3x: Pandas Dataframe of the dataset
    :param models_to_test: Number of Neural Networks to ensemble together
    :param lim1: variable to change architecture of Neural Network
    :param lim2: variable to change architecture of Neural Network
    :param lim3: variable to change architecture of Neural Network
    :param base_epochs: base number of epochs for the neural networks
    :param base_learning_rate: base learning rate of neural network
    :return: path to saved Neural Networks


    """

    print_and_write_status("Data Set Preparing for Neural Network Testing")
    print_and_write_status("------------------------------------------------------------")

    # Separating Training Data from Test Data

    Train_X = Whole_ETF_3x.sample(frac=0.90, random_state=0)
    Test_X = Whole_ETF_3x.drop(Train_X.index)

    train_features = Train_X.copy()
    test_features = Test_X.copy()
    prediction_features = prediction_ETF.copy()

    train_labels = train_features.pop('Profit Prediction' + str(future_num_days) + ' days from now')
    test_labels = test_features.pop('Profit Prediction' + str(future_num_days) + ' days from now')

    train_features = np.asarray(train_features).astype(float)
    test_features = np.asarray(test_features).astype(float)
    prediction_features = np.asarray(prediction_features).astype(float)

    train_labels = np.asarray(train_labels).astype(float)
    test_labels = np.asarray(test_labels).astype(float)

    # Normalize for input layer of Neural Network

    normalizer = tf.keras.layers.Normalization(axis=-1)
    normalizer.adapt(np.array(train_features))

    print_and_write_status("The Data has been prepared, the Neural Network is being Created")
    print_and_write_status("------------------------------------------------------------")

    ####################################################################################################################

    # Initialize Model(s)

    ####################################################################################################################

    # Input the number of Neural Networks to Ensemble together, initialize Test Model

    # models_to_test = 3
    test_model = None

    ####################################################################################################################

    # Begin Running Training for each model

    # Some initial details:

    #       - This is a feed-forward Neural Network with an Adam Loss Function
    #       - This neural network has 3 layers, each of which is different based on which number NN it is in the ensemble
    #       - The learning rate is changed based on which NN it is in the ensemble
    #       - The number of epochs is changed based on which NN it is in the ensemble
    #       - There is a slight learning decay which can be tuned if desired
    #       - It uses mean squared error to more aggressively deal with errors

    ####################################################################################################################

    for model_num in range(models_to_test):

        # If statement to run different kinds of Neural Networks - in this configuration, after NN #15, it will
        # Change from a 3 Hidden-Layer Neural Network to a 2 Hidden-Layer Network. This is an attempt to add
        # Variation to the Networks to help boost the outcome when they are averaged together.

        if model_num < lim1:

            test_model = keras.Sequential([
                normalizer,
                layers.Dense(model_num + 10, activation='relu'),
                layers.Dense(model_num + 10, activation='relu'),
                layers.Dense(model_num + 10, activation='relu'),
                layers.Dense(1)
            ])

        elif model_num < lim2:

            test_model = keras.Sequential([
                normalizer,
                layers.Dense(model_num, activation='relu'),
                layers.Dense(model_num, activation='relu'),
                layers.Dense(1)
            ])

        elif model_num < lim3:

            test_model = keras.Sequential([
                normalizer,
                layers.Dense(model_num - 5, activation='relu'),
                layers.Dense(model_num - 5, activation='relu'),
                layers.Dense(model_num - 5, activation='relu'),
                layers.Dense(1)
            ])

        else:

            test_model = keras.Sequential([
                normalizer,
                layers.Dense(model_num - 5, activation='relu'),
                layers.Dense(model_num - 5, activation='relu'),
                layers.Dense(1)
            ])

        # Help from https://keras.io/api/optimizers/learning_rate_schedules/exponential_decay/
        # Learning rate will decline as the number of epochs increases

        initial_learning_rate = base_learning_rate
        lr_schedule = tf.keras.optimizers.schedules.ExponentialDecay(
            initial_learning_rate,
            decay_steps=100000,
            decay_rate=0.97,
            staircase=True)

        test_model.compile(loss='mean_squared_error',
                           optimizer=tf.keras.optimizers.Adam(learning_rate=lr_schedule), metrics=['mse'])

        sys.stdout = open("../Website-GUI/status.txt", 'a')

        print("\n\nTraining: " + str(name_to_return) + '\n')

        print("The Neural Network ID Number: " + str(model_num) + " has been Created and Compiled")
        print("------------------------------------------------------------")

        print(
            "The Neural Network ID Number: " + str(model_num) + " is Starting to Run Over the Training Data")
        print("------------------------------------------------------------")

        history = test_model.fit(
            train_features,
            train_labels,
            validation_split=0.20,
            verbose=2, epochs=base_epochs)

        print("Saving Neural Network ID Number: " + str(model_num))
        print("------------------------------------------------------------")

        # Save Neural Network in Folder for 1) Later Reference, or 2) Manipulation from GUI

        test_model.save(
            '../Model-Training/Trained-Models/' + str(tfd['Model_Name']) + '/' + name_to_return + ' ' + str(
                model_num))

    ####################################################################################################################

    # Create Ensemble Prediction

    ####################################################################################################################

    # Help from https://medium.com/randomai/ensemble-and-store-models-in-keras-2-x-b881a6d7693f

    print("All Neural Networks Have Been Saved, Starting Ensemble Evaluation")
    print("------------------------------------------------------------")

    models = []
    for i in range(models_to_test):
        modelTemp = load_model(
            '../Model-Training/Trained-Models/' + str(tfd['Model_Name']) + '/' + name_to_return + ' ' + str(i))
        models.append(modelTemp)

    mean_ens_guess = []
    mean_predict_guess = []

    # Other options for choosing how to combine ensembled models

    # mode_ens_guess = []
    # max_ens_guess = []

    for point in test_features:
        raw_guess = []
        for model in models:
            # print_and_write_status(point)
            # print_and_write_status(len(point))
            # print_and_write_status(model)
            raw_guess.append(model.predict(point))

        # print_and_write_status(raw_guess)
        mean_ens_guess.append(np.mean(raw_guess))

    for point in prediction_features:
        # print_and_write_status(point)
        raw_guess = []
        for model in models:
            # print_and_write_status(point)
            # print_and_write_status(len(point))
            # print_and_write_status(model)
            raw_guess.append(model.predict(point))

        # print_and_write_status(raw_guess)
        mean_predict_guess.append(np.mean(raw_guess))

        # Other options for choosing how to combine ensembled models

        # mode_ens_guess.append(np.mode(raw_guess))
        # max_ens_guess.append(np.argmax(raw_guess)) - classification

    # y_output = [model.predict(test_features) for model in models]
    # y_average = layers.average(y_output)

    ####################################################################################################################

    # Plot Results of Predictions

    ####################################################################################################################

    print("The Neural Network is Starting to Run over the Test Data")
    print("------------------------------------------------------------")
    # guess = test_model.predict(test_features)

    train_loss = history.history['loss']
    val_loss = history.history['val_loss']

    plt.clf()

    print("Plotting Training Results and Saving to File")
    # print("------------------------------------------------------------")

    # Commented Out for Testing

    plt.subplot(1, 3, 1)
    plt.plot(history.history['loss'], label='loss')
    plt.plot(history.history['val_loss'], label='val_loss')
    plt.xlabel('Epoch')
    plt.ylabel('Error [Guess]')
    plt.legend()
    plt.grid(True)

    # Ensembled Neural Network Results

    plt.subplot(1, 3, 2)
    x = tf.linspace(0.0, len(test_labels) - 1, len(test_labels))
    plt.plot(x, test_labels, label='How Much Actually Traded')
    plt.plot(x, mean_ens_guess, label='Ensembled Network Guess at Price')
    plt.legend()

    # Ensembled Predictive Guess

    plt.subplot(1, 3, 3)
    x = tf.linspace(0.0, 30 - 1, 30)
    plt.plot(x, mean_predict_guess, label='Ensembled Network Guess at Future 30 day performance')
    plt.legend()

    # Save the Plot
    out_dir = "../Website-GUI/static/Training_Plots/" + str(tfd['Model_Name'])
    if not os.path.exists(out_dir):
        os.makedirs(out_dir)
    plt.savefig('../Website-GUI/static/Training_Plots/' + str(
        tfd['Model_Name']) + '/' + name_to_return + '_training_result.png')
    # plt.show()

    ####################################################################################################################

    # Storing results for later

    # test_results = {'test_model': test_model.evaluate(test_features, test_labels, verbose=0)}
    ####################################################################################################################
    return models_to_test, future_num_days, name_to_return, mean_predict_guess, train_loss, val_loss
    sys.stdout.close()


def EvaluateModels(list_of_lists, model_name_lists, train_loss_list, val_loss_list, model_name):
    plt.clf()
    flag = 0
    result_dict = {}
    for result in range(len(list_of_lists)):

        x = tf.linspace(0.0, 30 - 1, 30)
        plt.plot(x, list_of_lists[result], label=model_name_lists[result])

        result_dict[model_name_lists[result]] = [np.mean(list_of_lists[result]), np.mean(train_loss_list[result]),
                                                 np.mean(val_loss_list[result])]
        if flag == 0:
            max = np.mean(list_of_lists[result])
            max_name = model_name_lists[result]
            flag = 1
        else:
            if np.mean(list_of_lists[result]) > max:
                max = np.mean(list_of_lists[result])
                max_name = model_name_lists[result]
            else:
                pass
    plt.legend()
    plt.title("Predicted Performance Comparison")
    plt.xlabel('Days from Now')
    plt.ylabel('Percentage Predicted to be Profited')
    plt.savefig('./static/images/' + model_name + '.png')
    # plt.show()

    results_doc = docx.Document()
    results_doc.add_heading('Algorithm Projected Performance', 1)
    results_doc.add_heading('Results Plot', 3)
    results_doc.add_picture('./static/images/' + model_name + '.png')

    print_and_write_status(
        "The Best Overall Performing ETF Was: " + str(max_name) + ", with an averaged profit of: " + str(
            max) + "% over the next 30 days")

    results_doc.add_paragraph(
        "The Best Overall Performing ETF Was: " + str(max_name) + ", with an average growth of: " + str(max))

    print_and_write_status("Here are the results of the ETF's: ")
    for item in result_dict:
        print_and_write_status("ETF: " + item + " profited " + str(
            result_dict[item][0]) + "%. The Neural Network generated a Training Loss of ~" + str(
            int(result_dict[item][1])) + ", and a Validation Loss of ~" + str(int(result_dict[item][2])))
        results_doc.add_paragraph("ETF: " + item + " grew " + str(
            result_dict[item][0]) + "%. During testing, the Neural Network generated a Training Loss of ~" + str(
            int(result_dict[item][1])) + ", and a Validation Loss of ~" + str(int(result_dict[item][2])))

    out_dir2 = "./static/reports/" + str(tfd['Model_Name'])
    if not os.path.exists(out_dir2):
        os.makedirs(out_dir2)
    results_doc.save('./static/reports' + str(tfd['Model_Name']) + '/' + 'Comparison_Results.docx')
    convert("./staticComparison_Results.docx", str(tfd['Model_Name']) + '.pdf')


if __name__ == "__main__":
    # Send all output to status file
    def print_and_write_status(string):
        sys.stdout = open("../Website-GUI/status.txt", 'a')
        print(string)
        sys.stdout.close()


    # Put the training config into a dictionary
    with open("../Model-Training/training_config.txt", "r") as training_config_file:
        lines = training_config_file.readlines()

    # tfd stands for 'training config dict'
    tfd = {}
    for line in lines:
        for i in range(len(line)):
            if line[i] == ":":
                key = line[:i]
                value = line[(i + 2):-1]
                tfd[key] = value

    # Here is how I called the functions:
    '''
    ETF_created, num_days_to_predict, model_save_loc, prediction_ETF = CreateDataset()
    ETF_created2, num_days_to_predict2, model_save_loc2, prediction_ETF2 = CreateDataset('3x-ETF/SQQQ')
    ETF_created3, num_days_to_predict3, model_save_loc3, prediction_ETF3 = CreateDataset('3x-ETF/CURE')
    ETF_created4, num_days_to_predict4, model_save_loc4, prediction_ETF4 = CreateDataset('3x-ETF/SDOW')
    ETF_created5, num_days_to_predict5, model_save_loc5, prediction_ETF5 = CreateDataset('3x-ETF/SPXU')
    models_tested, num_days_predicted, name, results, train_loss, val_loss = CreateNeuralNetwork(ETF_created, prediction_ETF, num_days_to_predict,model_save_loc)
    models_tested2, num_days_predicted2, name2, results2, train_loss2, val_loss2 = CreateNeuralNetwork(ETF_created2, prediction_ETF2, num_days_to_predict2, model_save_loc2)
    models_tested3, num_days_predicted3, name3, results3, train_loss3, val_loss3 = CreateNeuralNetwork(ETF_created3, prediction_ETF2,
                                                                                   num_days_to_predict3, model_save_loc3)
    models_tested4, num_days_predicted4, name4, results4, train_loss4, val_loss4 = CreateNeuralNetwork(ETF_created4, prediction_ETF4,
                                                                                  num_days_to_predict4, model_save_loc4)
    models_tested5, num_days_predicted5, name5, results5, train_loss5, val_loss5 = CreateNeuralNetwork(ETF_created5, prediction_ETF,
                                                                                 num_days_to_predict5, model_save_loc5)
'''
    csvFiles = glob.glob(os.path.join("../Data/3x-ETF/", "*.csv"))
    listed = []
    listed_name = []
    listed_train_loss = []
    listed_val_loss = []
    total = len(csvFiles)
    current = 0
    for csvFile in csvFiles:
        current += 1
        print_and_write_status(str(current) + '/' + str(total))
        csvFile = csvFile[15:-4]
        # print_and_write_status("CSV FILE IS: " + csvFile)

        ETF_created, future_num_days, name_to_return, prediction_ETF = CreateDataset(name_to_return=str(csvFile),
                                                                                     lead_up_days=int(str(
                                                                                         tfd["Lead_Up_Days"]).replace(
                                                                                         " ", "")),
                                                                                     consideration_days=int(
                                                                                         tfd["Momentum_Consideration"]),
                                                                                     putcallflag=int(tfd["Putt_Call"]),
                                                                                     junkbondflag=int(tfd["Junk_Bond"]),
                                                                                     mcclellanflag=int(
                                                                                         tfd["McClellan_Summation"]))

        models_tested, num_days_predicted, name, results, train_loss, val_loss = CreateNeuralNetwork(
            Whole_ETF_3x=ETF_created, prediction_ETF=prediction_ETF,
            name_to_return=name_to_return, future_num_days=future_num_days,
            models_to_test=int(tfd["Models_to_Test"]),
            lim1=int(tfd["Limit1"]),
            lim2=int(tfd["Limit2"]),
            lim3=int(tfd["Limit3"]),
            base_epochs=int(tfd["Number_of_Epochs"]),
            base_learning_rate=float(tfd["Base_Learning_Rate"]))

        listed.append(results)
        listed_name.append(name)
        listed_train_loss.append(train_loss)
        listed_val_loss.append(val_loss)

    '''
    listed.append(results2)
    listed_name.append(name2)
    listed_train_loss.append(train_loss2)
    listed_val_loss.append(val_loss2)
    listed.append(results3)
    listed_train_loss.append(train_loss3)
    listed_val_loss.append(val_loss3)
    listed.append(results4)
    listed_name.append(name4)
    listed_train_loss.append(train_loss4)
    listed_val_loss.append(val_loss4)
    listed.append(results5)
    listed_name.append(name5)
    listed_train_loss.append(train_loss5)
    listed_val_loss.append(val_loss5)
    '''
    print_and_write_status("Evaluating Results...(This may take a while)")
    EvaluateModels(listed, listed_name, listed_train_loss, listed_val_loss, str(tfd["Model_Name"]))
    print_and_write_status("Cleaning Up...")

    # https://stackoverflow.com/questions/185936/how-to-delete-the-contents-of-a-folder
    files = glob.glob('../Data/Built-Datasets/*')
    for f in files:
        os.remove(f)

    print_and_write_status("Training Complete", end="")
