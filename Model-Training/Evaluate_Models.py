# Initial Network Testing
# Written by Nick Alvey

#####################################################################################################################
# Import Necessary Components
#####################################################################################################################

# For Data Parsing

# For Data Parsing

import pandas as pd
import glob
import os
import numpy as np
from scipy.stats import linregress
import datetime
import matplotlib.pyplot as plt

# For Neural Networks

import tensorflow as tf
from tensorflow import keras
from tensorflow.keras import layers
from keras.models import load_model

# For the Report

import docx
from docx2pdf import convert

# Add access to the data folder
sys.path.insert(1, '../Data')

def EvaluateModels(list_of_lists, model_name_lists, train_loss_list, val_loss_list):
    plt.clf()
    flag = 0
    result_dict = {}
    for result in range(len(list_of_lists)):

        x = tf.linspace(0.0, 30 - 1, 30)
        plt.plot(x, list_of_lists[result], label=model_name_lists[result])

        result_dict[model_name_lists[result]] = [np.mean(list_of_lists[result]), np.mean(train_loss_list[result]), np.mean(val_loss_list[result])]
        if flag == 0:
            max = np.mean(list_of_lists[result])
            max_name = model_name_lists[result]
            flag = 1
        else:
            if np.mean(list_of_lists[result]) > max:
                max = np.mean(list_of_lists[result])
                max_name = model_name_lists[result]
            else:
                pass
    plt.legend()
    plt.title("Predicted Performance Comparison")
    plt.xlabel('Days from Now')
    plt.ylabel('Percentage Predicted to be Profited')
    plt.savefig('Comparison_Plot.png')
    plt.show()

    results_doc = docx.Document()
    results_doc.add_heading('Algorithm Projected Performance', 1)
    results_doc.add_heading('Results Plot', 3)
    results_doc.add_picture('Comparison_Plot.png')


    print("The Best Overall Performing ETF Was: " + str(max_name) + ", with an averaged profit of: " + str(max) + "% over the next 30 days")

    results_doc.add_paragraph("The Best Overall Performing ETF Was: " + str(max_name) + ", with an average growth of: " + str(max))

    print("Here are the results of the ETF's: ")
    for item in result_dict:
        print("ETF: " + item + " profited " + str(result_dict[item][0]) + "%. The Neural Network generated a Training Loss of ~" + str(int(result_dict[item][1])) + ", and a Validation Loss of ~" + str(int(result_dict[item][2])))
        results_doc.add_paragraph("ETF: " + item + " grew " + str(result_dict[item][0]) + "%. During testing, the Neural Network generated a Training Loss of ~" + str(int(result_dict[item][1])) + ", and a Validation Loss of ~" + str(int(result_dict[item][2])))


    results_doc.save('Comparison_Results.docx')
    convert("Comparison_Results.docx")


if __name__ == "__main__":
    pass

# Here is how I called the functions:

# ETF_created,num_days_to_predict,model_save_loc, prediction_ETF = CreateDataset()
#     ETF_created2, num_days_to_predict2, model_save_loc2, prediction_ETF2 = CreateDataset('3x-ETF/SQQQ')
#     ETF_created3, num_days_to_predict3, model_save_loc3, prediction_ETF3 = CreateDataset('3x-ETF/CURE')
#     ETF_created4, num_days_to_predict4, model_save_loc4, prediction_ETF4 = CreateDataset('3x-ETF/SDOW')
#     ETF_created5, num_days_to_predict5, model_save_loc5, prediction_ETF5 = CreateDataset('3x-ETF/SPXU')
#     models_tested, num_days_predicted, name, results, train_loss, val_loss = CreateNeuralNetwork(ETF_created, prediction_ETF, num_days_to_predict,model_save_loc)
#     models_tested2, num_days_predicted2, name2, results2, train_loss2, val_loss2 = CreateNeuralNetwork(ETF_created2, prediction_ETF2, num_days_to_predict2, model_save_loc2)
#     models_tested3, num_days_predicted3, name3, results3, train_loss3, val_loss3 = CreateNeuralNetwork(ETF_created3, prediction_ETF2,
#                                                                                num_days_to_predict3, model_save_loc3)
#     models_tested4, num_days_predicted4, name4, results4, train_loss4, val_loss4 = CreateNeuralNetwork(ETF_created4, prediction_ETF4,
#                                                                                num_days_to_predict4, model_save_loc4)
#     models_tested5, num_days_predicted5, name5, results5, train_loss5, val_loss5 = CreateNeuralNetwork(ETF_created5, prediction_ETF,
#                                                                                num_days_to_predict5, model_save_loc5)
#
#     listed = []
#     listed_name = []
#     listed_train_loss = []
#     listed_val_loss = []
#
#     listed.append(results)
#     listed_name.append(name)
#     listed_train_loss.append(train_loss)
#     listed_val_loss.append(val_loss)
#     listed.append(results2)
#     listed_name.append(name2)
#     listed_train_loss.append(train_loss2)
#     listed_val_loss.append(val_loss2)
#     listed.append(results3)
#     listed_name.append(name3)
#     listed_train_loss.append(train_loss3)
#     listed_val_loss.append(val_loss3)
#     listed.append(results4)
#     listed_name.append(name4)
#     listed_train_loss.append(train_loss4)
#     listed_val_loss.append(val_loss4)
#     listed.append(results5)
#     listed_name.append(name5)
#     listed_train_loss.append(train_loss5)
#     listed_val_loss.append(val_loss5)
#
#
#     EvaluateModels(listed, listed_name, listed_train_loss, listed_val_loss)
